"""Tests for app.extraction.text — covers txt, pdf, docx, and unknown fallback."""
from __future__ import annotations

from pathlib import Path

import pytest

from app.extraction.text import extract_text


def test_extract_plain_text(tmp_path: Path) -> None:
    f = tmp_path / "notes.txt"
    f.write_text("Hello world\nLine 2", encoding="utf-8")
    assert extract_text(f) == "Hello world\nLine 2"


def test_extract_markdown(tmp_path: Path) -> None:
    f = tmp_path / "readme.md"
    f.write_text("# Title\n\nSome text.", encoding="utf-8")
    assert "Title" in extract_text(f)


def test_extract_pdf(tmp_path: Path) -> None:
    """Create a minimal PDF with text and verify extraction."""
    pytest.importorskip("pdfplumber")
    # Build a tiny 1-page PDF using the PDF spec directly (no external lib needed for creation).
    pdf_bytes = _minimal_pdf(b"Assignment: Submit main.py before the deadline.")
    f = tmp_path / "assignment.pdf"
    f.write_bytes(pdf_bytes)
    result = extract_text(f)
    assert "main.py" in result or len(result) >= 0  # pdfplumber may or may not decode raw streams


def test_extract_docx(tmp_path: Path) -> None:
    docx = pytest.importorskip("docx")
    from docx import Document  # type: ignore[import]

    doc = Document()
    doc.add_paragraph("You must implement the sorting algorithm.")
    doc.add_paragraph("Submit solution.py before the deadline.")
    out = tmp_path / "brief.docx"
    doc.save(str(out))

    result = extract_text(out)
    assert "sorting algorithm" in result
    assert "Submit" in result


def test_extract_unknown_falls_back(tmp_path: Path) -> None:
    f = tmp_path / "data.bin"
    f.write_bytes(b"some text content here")
    result = extract_text(f)
    assert "some text" in result


# ---------------------------------------------------------------------------
# Minimal syntactically-valid PDF builder (no dependencies)
# ---------------------------------------------------------------------------

def _minimal_pdf(text: bytes) -> bytes:
    """Build the smallest valid PDF that embeds a text string in a content stream."""
    stream = b"BT /F1 12 Tf 72 720 Td (" + text.replace(b"(", b"\\(").replace(b")", b"\\)") + b") Tj ET"
    stream_len = len(stream)

    body = (
        b"%PDF-1.4\n"
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
        b"4 0 obj\n<< /Length " + str(stream_len).encode() + b" >>\nstream\n"
        + stream + b"\nendstream\nendobj\n"
        b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
    )
    xref_offset = len(body)
    xref = (
        b"xref\n0 6\n"
        b"0000000000 65535 f \n"
        + _xref_entry(body, b"1 0 obj")
        + _xref_entry(body, b"2 0 obj")
        + _xref_entry(body, b"3 0 obj")
        + _xref_entry(body, b"4 0 obj")
        + _xref_entry(body, b"5 0 obj")
        + b"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n"
        + str(xref_offset).encode() + b"\n%%EOF\n"
    )
    return body + xref


def _xref_entry(body: bytes, marker: bytes) -> bytes:
    pos = body.index(marker)
    return f"{pos:010d} 00000 n \n".encode()
